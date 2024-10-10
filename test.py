import pandas as pd
import numpy as np
from datetime import datetime
import csv
import os
from docx import Document
from docx.shared import Inches

def analyze_dataset(file_path):
    # Read the CSV file
    df = pd.read_csv(file_path)
    
    # Dictionary to store field analysis results
    field_analysis = {}
    
    # Analyze each field
    for column in df.columns:
        field_analysis[column] = analyze_field(df[column])
    
    # Flag outliers
    flagged_records = flag_outliers(df, field_analysis)
    
    return field_analysis, flagged_records

def analyze_field(series):
    # Determine data type
    if pd.api.types.is_bool_dtype(series):
        data_type = 'boolean'
    elif pd.api.types.is_numeric_dtype(series):
        data_type = 'numeric'
    elif pd.api.types.is_datetime64_any_dtype(series):
        data_type = 'date'
    else:
        data_type = 'string'
    
    # Get basic statistics
    stats = {
        'data_type': data_type,
        'unique_values': int(series.nunique()),
        'missing_values': int(series.isnull().sum())
    }
    
    if data_type == 'numeric':
        stats.update({
            'mean': f"{series.mean():.2f}",
            'median': f"{series.median():.2f}",
            'std': f"{series.std():.2f}",
            'min': f"{series.min():.2f}",
            'max': f"{series.max():.2f}"
        })
    elif data_type == 'date':
        stats.update({
            'min_date': series.min().strftime('%Y-%m-%d'),
            'max_date': series.max().strftime('%Y-%m-%d')
        })
    elif data_type == 'boolean':
        stats.update({
            'true_count': int(series.sum()),
            'false_count': int((~series).sum())
        })
    
    return stats

def flag_outliers(df, field_analysis):
    flagged_records = []
    
    for column, analysis in field_analysis.items():
        if analysis['data_type'] == 'numeric':
            flagged_records.extend(flag_numeric_outliers(df, column, analysis))
        elif analysis['data_type'] == 'date':
            flagged_records.extend(flag_date_outliers(df, column, analysis))
        elif analysis['data_type'] == 'boolean':
            flagged_records.extend(flag_boolean_outliers(df, column, analysis))
        else:  # string
            flagged_records.extend(flag_string_outliers(df, column, analysis))
    
    return flagged_records

def flag_numeric_outliers(df, column, analysis):
    flagged = []
    series = df[column]
    
    # Check for negatives
    if series.min() >= 0:
        negatives = series[series < 0]
        for idx in negatives.index:
            flagged.append((idx, column, f"Negative value: {series[idx]}"))
    
    # Check for outliers using IQR method
    Q1 = series.quantile(0.25)
    Q3 = series.quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    
    outliers = series[(series < lower_bound) | (series > upper_bound)]
    for idx in outliers.index:
        flagged.append((idx, column, f"Outlier value: {series[idx]}"))
    
    return flagged

def flag_date_outliers(df, column, analysis):
    flagged = []
    series = pd.to_datetime(df[column])
    
    # Check for dates after the maximum allowed date
    max_allowed_date = pd.to_datetime('2024-09-30')
    future_dates = series[series > max_allowed_date]
    for idx in future_dates.index:
        flagged.append((idx, column, f"Future date: {series[idx]}"))
    
    # Check for unusually old dates (eset manually by analyst)
    median_date = series.median()
    old_threshold = median_date - pd.Timedelta(days=30000) 
    old_dates = series[series < old_threshold]
    for idx in old_dates.index:
        flagged.append((idx, column, f"Unusually old date: {series[idx]}"))
    
    return flagged

def flag_boolean_outliers(df, column, analysis):
    flagged = []
    series = df[column]
    
    # Check for any non-boolean values
    non_boolean = series[~series.isin([True, False, 0, 1])]
    for idx in non_boolean.index:
        flagged.append((idx, column, f"Non-boolean value: {series[idx]}"))
    
    return flagged

def flag_string_outliers(df, column, analysis):
    flagged = []
    series = df[column]
    
    # Check for missing values
    missing = series[series.isnull()]
    for idx in missing.index:
        flagged.append((idx, column, "Missing value"))
    
    # Check for placeholder values
    placeholders = ['TBD', 'N/A', 'Unknown']
    for placeholder in placeholders:
        matches = series[series.str.lower() == placeholder.lower()]
        for idx in matches.index:
            flagged.append((idx, column, f"Placeholder value: {series[idx]}"))
    
    # Check for unusually short or long values
    lengths = series.str.len()
    mean_length = lengths.mean()
    std_length = lengths.std()
    
    short_values = series[lengths < (mean_length - 2 * std_length)]
    for idx in short_values.index:
        flagged.append((idx, column, f"Unusually short value: {series[idx]}"))
    
    long_values = series[lengths > (mean_length + 2 * std_length)]
    for idx in long_values.index:
        flagged.append((idx, column, f"Unusually long value: {series[idx]}"))
    
    return flagged

def save_analysis_to_file(field_analysis, flagged_records, input_file_path):
    # Create output directory
    output_dir = os.path.join(os.path.dirname(input_file_path), 'analysis_output')
    os.makedirs(output_dir, exist_ok=True)
    
    # Save field analysis to DOCX file
    analysis_file_path = os.path.join(output_dir, r'C:\Store\Projects\IAA\2024 10 10 Data cleaning\field_analysis.docx')
    doc = Document()
    doc.add_heading('Field Analysis', 0)

    # Add a table for the field analysis
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field Name'
    hdr_cells[1].text = 'Data Type'
    hdr_cells[2].text = 'Unique Values'
    hdr_cells[3].text = 'Missing Values'
    hdr_cells[4].text = 'Mean'
    hdr_cells[5].text = 'Median'
    hdr_cells[6].text = 'Min'
    hdr_cells[7].text = 'Max'

    for field, analysis in field_analysis.items():
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = analysis['data_type']
        row_cells[2].text = str(analysis['unique_values'])
        row_cells[3].text = str(analysis['missing_values'])
        if analysis['data_type'] == 'numeric':
            row_cells[4].text = analysis['mean']
            row_cells[5].text = analysis['median']
            row_cells[6].text = analysis['min']
            row_cells[7].text = analysis['max']
        elif analysis['data_type'] == 'date':
            row_cells[6].text = analysis['min_date']
            row_cells[7].text = analysis['max_date']
        elif analysis['data_type'] == 'boolean':
            row_cells[4].text = f"True: {analysis['true_count']}"
            row_cells[5].text = f"False: {analysis['false_count']}"

    # Adjust column widths
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(1)

    doc.save(analysis_file_path)
    
    # Save flagged records to CSV file
    flagged_file_path = os.path.join(output_dir, r'C:\Store\Projects\IAA\2024 10 10 Data cleaning\flagged_records.csv')
    with open(flagged_file_path, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['Row', 'Column', 'Reason'])
        writer.writerows(flagged_records)
    
    return analysis_file_path, flagged_file_path

# Example usage
if __name__ == "__main__":
    file_path = r"C:\Store\Projects\IAA\2024 10 10 Data cleaning\insurance_dummy_data.csv"
    field_analysis, flagged_records = analyze_dataset(file_path)
    
    # Save results to files
    analysis_file, flagged_file = save_analysis_to_file(field_analysis, flagged_records, file_path)
    
    print(f"Field analysis saved to: {analysis_file}")
    print(f"Flagged records saved to: {flagged_file}")
    
    print("\nField Analysis Summary:")
    for field, analysis in field_analysis.items():
        print(f"{field}: {analysis['data_type']}, {analysis['unique_values']} unique values, {analysis['missing_values']} missing values")
    
    print(f"\nTotal flagged records: {len(flagged_records)}")